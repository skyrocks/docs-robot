from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import pymysql


dbName = 'xxxxxx'
dbUser = 'root'
dbPWd = 'root'
dbHost = 'localhost'
dbPort = 3306
# 数据库链接
db = None
title = '数据库文档'

# word表格列宽
col_width_dic = {0: 0.5, 1: 2, 2: 3.5, 3: 1, 4: 1, 5: 1, 6: 1}


def create_connect():
    global db
    db = pymysql.connect(host=dbHost, port=dbPort, user=dbUser, passwd=dbPWd, db=dbName, charset='utf8')


def get_table():
    if db is None:
        create_connect()
    cursor = db.cursor()
    sql = '''SELECT 
                table_name, TABLE_COMMENT 
            FROM information_schema.`TABLES` 
            WHERE TABLE_SCHEMA = '%s' 
            order by table_name;'''%(dbName)
    cursor.execute(sql)
    return cursor


def get_field(table_code):
    if db is None:
        create_connect()
    cursor = db.cursor()
    sql = '''SELECT
                COLUMN_NAME,
                ORDINAL_POSITION,
                case when COLUMN_DEFAULT is null or COLUMN_DEFAULT = 'NULL' then '' else COLUMN_DEFAULT end,
                case when IS_NULLABLE = 'NO' then '' else '空' end,
                COLUMN_TYPE,
                case when COLUMN_KEY = 'PRI' then '是' else '' end,
                COLUMN_COMMENT
            FROM
                information_schema.`COLUMNS`
            WHERE
                TABLE_SCHEMA = 'parrot' and TABLE_NAME = '%s'
            ORDER BY
                TABLE_NAME,
                ORDINAL_POSITION;'''%(table_code)
    cursor.execute(sql)
    return cursor


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def create_doc():

    document = Document()
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    for i in range(5):
        document.add_paragraph()

    # 添加标题，并设置级别，范围：0 至 9，默认为1
    h = document.add_heading(title, 0)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(10):
        document.add_paragraph()

    p = document.add_paragraph(datetime.datetime.now().strftime('%Y-%m-%d'))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    table_count = 0
    cursor = get_table()
    while 1:
        res = cursor.fetchone()
        if res is None:
            break

        table_count = table_count + 1
        table_code = res[0]
        table_name = res[1]

        head = document.add_heading(level=1)
        head_run = head.add_run(str(table_count)+ '\t'+table_name + "  " + table_code)
        head_run.font.color.rgb = RGBColor(0, 0, 0)  # 字体颜色
        #t.bold = True
        #t.style.font.size = Pt(14)

        table = document.add_table(rows=1, cols=7, style='Light List Accent 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '字段'
        hdr_cells[2].text = '名称'
        hdr_cells[3].text = '类型'
        hdr_cells[4].text = '默认值'
        hdr_cells[5].text = '主键'
        hdr_cells[6].text = '允许空值'

        for col in hdr_cells:
            set_cell_border(col, end={"val": "single"})
        for col_num in range(7):
            hdr_cells[col_num].width = Inches(col_width_dic[col_num])

        records = get_field(table_code)
        while 1:
            rec = records.fetchone()
            if rec is None:
                break

            field_code = rec[0]
            field_seq = rec[1]
            field_default = rec[2]
            field_null = rec[3]
            field_type = rec[4]
            field_pri = rec[5]
            field_name = rec[6]

            # 表格添加行，并返回行所在的单元格列表
            row_cells = table.add_row().cells
            row_cells[0].text = str(field_seq)
            row_cells[1].text = field_code
            row_cells[2].text = field_name
            row_cells[3].text = field_type
            row_cells[4].text = field_default
            row_cells[5].text = field_pri
            row_cells[6].text = field_null

            for c in row_cells:
                set_cell_border(c, end={"val": "single"})
            for col_num in range(7):
                row_cells[col_num].width = Inches(col_width_dic[col_num])

        records.close()

        document.add_paragraph()

    # 关闭游标
    cursor.close()

    db.close()

    # 保存.docx文档
    document.save(title+dbName+'('+datetime.datetime.now().strftime('%Y%m%d')+').docx')


if __name__ == '__main__':
    create_doc()
